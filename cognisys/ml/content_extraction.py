"""
Content Extraction Layer for CogniSys

Extracts text content from various file formats to enable content-based classification.
Supports PDFs, Word documents, Excel spreadsheets, images (OCR), and text files.

Author: Claude Code
Date: 2025-12-01
"""

import logging
from pathlib import Path
from typing import Dict, Optional, Tuple
import mimetypes

logger = logging.getLogger(__name__)


class ContentExtractor:
    """
    Extract text content from various file formats.

    Supports:
    - PDFs (PyMuPDF)
    - Word documents (python-docx)
    - Excel spreadsheets (openpyxl)
    - Images (NVIDIA Vision OCR)
    - Text files (direct read)
    - Code files (direct read)
    """

    def __init__(self, max_chars: int = 2000, nvidia_api_key: Optional[str] = None):
        """
        Initialize content extractor.

        Args:
            max_chars: Maximum characters to extract per file
            nvidia_api_key: Optional NVIDIA API key for OCR
        """
        self.max_chars = max_chars
        self.nvidia_api_key = nvidia_api_key

        # Check available extractors
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which extraction libraries are available"""
        self.has_pymupdf = False
        self.has_docx = False
        self.has_openpyxl = False
        self.has_pil = False
        self.has_nvidia = bool(self.nvidia_api_key)

        try:
            import fitz  # PyMuPDF
            self.has_pymupdf = True
        except ImportError:
            logger.warning("PyMuPDF not installed - PDF extraction disabled")

        try:
            import docx
            self.has_docx = True
        except ImportError:
            logger.warning("python-docx not installed - Word extraction disabled")

        try:
            import openpyxl
            self.has_openpyxl = True
        except ImportError:
            logger.warning("openpyxl not installed - Excel extraction disabled")

        try:
            from PIL import Image
            self.has_pil = True
        except ImportError:
            logger.warning("PIL not installed - Image extraction disabled")

    def extract(self, file_path: Path) -> Dict[str, any]:
        """
        Extract content from file.

        Args:
            file_path: Path to file

        Returns:
            Dictionary with:
            - content: Extracted text (first max_chars characters)
            - full_content: Full extracted text (if needed)
            - method: Extraction method used
            - success: Boolean indicating success
            - error: Error message if failed
            - metadata: Additional metadata (page count, word count, etc.)
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return {
                'content': '',
                'method': None,
                'success': False,
                'error': f'File not found: {file_path}'
            }

        # Determine file type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        extension = file_path.suffix.lower()

        # Route to appropriate extractor
        try:
            if extension == '.pdf':
                return self._extract_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return self._extract_word(file_path)
            elif extension in ['.xlsx', '.xls']:
                return self._extract_excel(file_path)
            elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                return self._extract_image(file_path)
            elif extension in ['.txt', '.md', '.csv', '.log', '.json', '.xml', '.yaml', '.yml']:
                return self._extract_text(file_path)
            elif extension in ['.py', '.js', '.java', '.cpp', '.c', '.h', '.cs', '.rb', '.go', '.rs']:
                return self._extract_text(file_path)
            else:
                # Try text extraction as fallback
                return self._extract_text(file_path)

        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return {
                'content': '',
                'method': 'error',
                'success': False,
                'error': str(e)
            }

    def _extract_pdf(self, file_path: Path) -> Dict:
        """Extract text from PDF using PyMuPDF"""
        if not self.has_pymupdf:
            return {
                'content': '',
                'method': 'pdf_unavailable',
                'success': False,
                'error': 'PyMuPDF not installed'
            }

        import fitz  # PyMuPDF

        try:
            doc = fitz.open(str(file_path))

            # Extract text from all pages
            full_text = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                full_text.append(page.get_text())

            full_content = '\n'.join(full_text)
            content = full_content[:self.max_chars]

            # Metadata
            metadata = {
                'page_count': len(doc),
                'char_count': len(full_content),
                'word_count': len(full_content.split())
            }

            doc.close()

            return {
                'content': content,
                'full_content': full_content,
                'method': 'pymupdf',
                'success': True,
                'metadata': metadata
            }

        except Exception as e:
            return {
                'content': '',
                'method': 'pymupdf_error',
                'success': False,
                'error': str(e)
            }

    def _extract_word(self, file_path: Path) -> Dict:
        """Extract text from Word document using python-docx"""
        if not self.has_docx:
            return {
                'content': '',
                'method': 'docx_unavailable',
                'success': False,
                'error': 'python-docx not installed'
            }

        import docx

        try:
            doc = docx.Document(str(file_path))

            # Extract text from all paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            full_content = '\n'.join(full_text)
            content = full_content[:self.max_chars]

            # Metadata
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'char_count': len(full_content),
                'word_count': len(full_content.split())
            }

            return {
                'content': content,
                'full_content': full_content,
                'method': 'python-docx',
                'success': True,
                'metadata': metadata
            }

        except Exception as e:
            return {
                'content': '',
                'method': 'docx_error',
                'success': False,
                'error': str(e)
            }

    def _extract_excel(self, file_path: Path) -> Dict:
        """Extract text from Excel spreadsheet using openpyxl"""
        if not self.has_openpyxl:
            return {
                'content': '',
                'method': 'excel_unavailable',
                'success': False,
                'error': 'openpyxl not installed'
            }

        import openpyxl

        try:
            workbook = openpyxl.load_workbook(str(file_path), data_only=True)

            # Extract text from all sheets
            full_text = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                full_text.append(f"Sheet: {sheet_name}")

                # Extract cell values (limit to first 100 rows to avoid huge sheets)
                for row_idx, row in enumerate(sheet.iter_rows(max_row=100, values_only=True)):
                    row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        full_text.append(row_text)

            full_content = '\n'.join(full_text)
            content = full_content[:self.max_chars]

            # Metadata
            metadata = {
                'sheet_count': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames,
                'char_count': len(full_content),
                'word_count': len(full_content.split())
            }

            return {
                'content': content,
                'full_content': full_content,
                'method': 'openpyxl',
                'success': True,
                'metadata': metadata
            }

        except Exception as e:
            return {
                'content': '',
                'method': 'excel_error',
                'success': False,
                'error': str(e)
            }

    def _extract_image(self, file_path: Path) -> Dict:
        """Extract text from image using NVIDIA Vision OCR (if available)"""
        if not self.has_nvidia:
            return {
                'content': '',
                'method': 'ocr_unavailable',
                'success': False,
                'error': 'NVIDIA API key not provided - OCR disabled'
            }

        # TODO: Implement NVIDIA Vision OCR
        # This will be implemented in the next step
        return {
            'content': '',
            'method': 'nvidia_ocr_pending',
            'success': False,
            'error': 'NVIDIA Vision OCR not yet implemented'
        }

    def _extract_text(self, file_path: Path) -> Dict:
        """Extract text from plain text file"""
        try:
            # Try common encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        full_content = f.read()

                    content = full_content[:self.max_chars]

                    # Metadata
                    metadata = {
                        'encoding': encoding,
                        'char_count': len(full_content),
                        'word_count': len(full_content.split()),
                        'line_count': len(full_content.splitlines())
                    }

                    return {
                        'content': content,
                        'full_content': full_content,
                        'method': 'text_direct',
                        'success': True,
                        'metadata': metadata
                    }

                except (UnicodeDecodeError, UnicodeError):
                    continue

            # If all encodings fail, try binary
            with open(file_path, 'rb') as f:
                binary_content = f.read()

            # Try to decode as best as possible
            content = binary_content[:self.max_chars].decode('utf-8', errors='ignore')

            return {
                'content': content,
                'method': 'text_binary',
                'success': True,
                'metadata': {'encoding': 'binary'}
            }

        except Exception as e:
            return {
                'content': '',
                'method': 'text_error',
                'success': False,
                'error': str(e)
            }


def extract_content_from_file(file_path: str, max_chars: int = 2000) -> Tuple[str, str]:
    """
    Convenience function to extract content from a file.

    Args:
        file_path: Path to file
        max_chars: Maximum characters to extract

    Returns:
        Tuple of (content, extraction_method)
    """
    extractor = ContentExtractor(max_chars=max_chars)
    result = extractor.extract(Path(file_path))

    return result.get('content', ''), result.get('method', 'unknown')


# Example usage
if __name__ == '__main__':
    import sys

    logging.basicConfig(level=logging.INFO)

    if len(sys.argv) < 2:
        print("Usage: python content_extraction.py <file_path>")
        sys.exit(1)

    file_path = sys.argv[1]

    extractor = ContentExtractor(max_chars=500)
    result = extractor.extract(Path(file_path))

    print(f"\n{'='*80}")
    print(f"CONTENT EXTRACTION RESULT")
    print(f"{'='*80}")
    print(f"File: {file_path}")
    print(f"Method: {result['method']}")
    print(f"Success: {result['success']}")

    if result['success']:
        print(f"\nContent (first 500 chars):")
        print(f"{'-'*80}")
        print(result['content'])
        print(f"{'-'*80}")

        if 'metadata' in result:
            print(f"\nMetadata:")
            for key, value in result['metadata'].items():
                print(f"  {key}: {value}")
    else:
        print(f"\nError: {result.get('error', 'Unknown error')}")

    print(f"{'='*80}\n")
